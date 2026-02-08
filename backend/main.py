import os
import shutil
import tempfile
import uuid
from typing import List, Optional
from datetime import datetime, timedelta
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks, Depends, status
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from jose import JWTError, jwt
from passlib.context import CryptContext
from pydantic import BaseModel

# Standard Conversion Libraries (No AI/ML)
from pdf2docx import Converter as PDFConverter
from docx2pdf import convert as docx_to_pdf
import pandas as pd
from PIL import Image
import img2pdf
import fitz  # PyMuPDF
from fpdf import FPDF
import pdfplumber
from docx import Document

# --- SECURITY CONFIG ---
SECRET_KEY = "your-secret-key-keep-it-secret" # In production, use environment variable
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 30

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")

# --- MOCK DATABASE (Use SQLite/PostgreSQL for production) ---
users_db = {}

class User(BaseModel):
    username: str
    email: Optional[str] = None
    disabled: Optional[bool] = None

class UserInDB(User):
    hashed_password: str

class Token(BaseModel):
    access_token: str
    token_type: str

class UserCreate(BaseModel):
    username: str
    password: str
    email: str

def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password):
    return pwd_context.hash(password)

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

async def get_current_user(token: str = Depends(oauth2_scheme)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise credentials_exception
    except JWTError:
        raise credentials_exception
    user = users_db.get(username)
    if user is None:
        raise credentials_exception
    return user

# --- APP SETUP ---
app = FastAPI(title="Type Shift API")

# Enable CORS so the React frontend can talk to this API
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

TEMP_DIR = tempfile.gettempdir()

# --- AUTH ENDPOINTS ---

@app.post("/register")
async def register(user: UserCreate):
    if user.username in users_db:
        raise HTTPException(status_code=400, detail="Username already registered")
    hashed_password = get_password_hash(user.password)
    users_db[user.username] = {
        "username": user.username,
        "email": user.email,
        "hashed_password": hashed_password,
        "disabled": False
    }
    return {"message": "User created successfully"}

@app.post("/token", response_model=Token)
async def login_for_access_token(form_data: OAuth2PasswordRequestForm = Depends()):
    user_dict = users_db.get(form_data.username)
    if not user_dict or not verify_password(form_data.password, user_dict["hashed_password"]):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user_dict["username"]}, expires_delta=access_token_expires
    )
    return {"access_token": access_token, "token_type": "bearer"}

@app.get("/users/me", response_model=User)
async def read_users_me(current_user: User = Depends(get_current_user)):
    return current_user

@app.get("/")
async def root():
    """Health check endpoint to verify server is running."""
    return {"status": "online", "message": "SwiftConvert API is running"}

def cleanup_files(file_paths: List[str]):
    """Removes temporary files from the server after the response is sent."""
    for path in file_paths:
        try:
            if os.path.exists(path):
                os.remove(path)
                print(f"Cleaned up: {path}")
        except Exception as e:
            print(f"Error deleting file {path}: {e}")

@app.post("/convert")
async def convert_file(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    target_format: str = Form(...)
):
    """
    Main conversion engine. 
    1. Saves the uploaded file to a temporary location.
    2. Runs the appropriate library for conversion based on input/output types.
    3. Schedules file cleanup as a background task.
    """
    original_ext = file.filename.split(".")[-1].lower()
    target_format = target_format.upper()
    
    # Generate unique filenames to avoid collisions
    input_filename = f"{uuid.uuid4()}.{original_ext}"
    input_path = os.path.join(TEMP_DIR, input_filename)

    # Save the uploaded file locally
    with open(input_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    output_filename = str(uuid.uuid4())
    output_path = ""
    
    try:
        # --- DOCUMENT CONVERSIONS ---
        if original_ext == "pdf" and target_format == "DOCX":
            output_path = os.path.join(TEMP_DIR, f"{output_filename}.docx")
            cv = PDFConverter(input_path)
            cv.convert(output_path) # Uses algorithmic structural analysis (non-AI) to reconstruct Word docs
            cv.close()

        elif original_ext == "docx" and target_format == "PDF":
            output_path = os.path.join(TEMP_DIR, f"{output_filename}.pdf")
            docx_to_pdf(input_path, output_path)

        # --- NEW: PDF TO EXCEL/CSV ---
        elif original_ext == "pdf" and target_format in ["XLSX", "CSV"]:
            all_tables = []
            with pdfplumber.open(input_path) as pdf:
                for page in pdf.pages:
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            df = pd.DataFrame(table[1:], columns=table[0])
                            all_tables.append(df)
            
            if not all_tables:
                # If no tables found, extract all text as a single row per line
                with pdfplumber.open(input_path) as pdf:
                    text = ""
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                    all_tables.append(pd.DataFrame([line.split() for line in text.split('\n') if line.strip()]))

            final_df = pd.concat(all_tables, ignore_index=True) if all_tables else pd.DataFrame()
            
            if target_format == "XLSX":
                output_path = os.path.join(TEMP_DIR, f"{output_filename}.xlsx")
                final_df.to_excel(output_path, index=False)
            else:
                output_path = os.path.join(TEMP_DIR, f"{output_filename}.csv")
                final_df.to_csv(output_path, index=False)

        # --- NEW: DOCX TO EXCEL/CSV ---
        elif original_ext == "docx" and target_format in ["XLSX", "CSV"]:
            doc = Document(input_path)
            all_tables = []
            for table in doc.tables:
                data = []
                for row in table.rows:
                    data.append([cell.text.strip() for cell in row.cells])
                if data:
                    df = pd.DataFrame(data[1:], columns=data[0]) if len(data) > 1 else pd.DataFrame(data)
                    all_tables.append(df)

            if not all_tables:
                # Extract paragraph text if no tables
                text = [p.text for p in doc.paragraphs if p.text.strip()]
                all_tables.append(pd.DataFrame({"Content": text}))

            final_df = pd.concat(all_tables, ignore_index=True) if all_tables else pd.DataFrame()

            if target_format == "XLSX":
                output_path = os.path.join(TEMP_DIR, f"{output_filename}.xlsx")
                final_df.to_excel(output_path, index=False)
            else:
                output_path = os.path.join(TEMP_DIR, f"{output_filename}.csv")
                final_df.to_csv(output_path, index=False)

        # --- SPREADSHEET CONVERSIONS ---
        elif original_ext == "csv" and target_format == "XLSX":
            output_path = os.path.join(TEMP_DIR, f"{output_filename}.xlsx")
            pd.read_csv(input_path).to_excel(output_path, index=False)
            
        elif original_ext == "xlsx" and target_format == "CSV":
            output_path = os.path.join(TEMP_DIR, f"{output_filename}.csv")
            pd.read_excel(input_path).to_csv(output_path, index=False)

        elif original_ext in ["csv", "xlsx"] and target_format == "PDF":
            output_path = os.path.join(TEMP_DIR, f"{output_filename}.pdf")
            df = pd.read_csv(input_path) if original_ext == "csv" else pd.read_excel(input_path)
            
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("helvetica", size=10)
            
            cols = df.columns.tolist()
            for col in cols:
                pdf.cell(35, 10, str(col)[:15], border=1)
            pdf.ln()
            for i, row in df.head(100).iterrows():
                for val in row:
                    pdf.cell(35, 10, str(val)[:15], border=1)
                pdf.ln()
            pdf.output(output_path)

        # --- IMAGE CONVERSIONS ---
        elif original_ext in ["jpg", "jpeg", "png"] and target_format == "PDF":
            output_path = os.path.join(TEMP_DIR, f"{output_filename}.pdf")
            with open(output_path, "wb") as f:
                f.write(img2pdf.convert(input_path))

        elif original_ext in ["jpg", "jpeg"] and target_format == "PNG":
            output_path = os.path.join(TEMP_DIR, f"{output_filename}.png")
            img = Image.open(input_path)
            img.save(output_path, "PNG")

        elif original_ext == "png" and target_format == "JPG":
            output_path = os.path.join(TEMP_DIR, f"{output_filename}.jpg")
            img = Image.open(input_path).convert("RGB")
            img.save(output_path, "JPEG")

        # --- PDF TO IMAGE ---
        elif original_ext == "pdf" and target_format in ["PNG", "JPG"]:
            ext = target_format.lower()
            output_path = os.path.join(TEMP_DIR, f"{output_filename}.{ext}")
            doc = fitz.open(input_path)
            page = doc.load_page(0)
            pix = page.get_pixmap()
            pix.save(output_path)
            doc.close()

        else:
            raise HTTPException(status_code=400, detail=f"Unsupported conversion: {original_ext} to {target_format}")

        background_tasks.add_task(cleanup_files, [input_path, output_path])
        
        return FileResponse(
            path=output_path,
            filename=f"converted_{file.filename.split('.')[0]}.{target_format.lower()}",
            media_type='application/octet-stream'
        )

    except Exception as e:
        cleanup_files([input_path])
        if output_path: cleanup_files([output_path])
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)
